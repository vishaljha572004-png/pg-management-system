import os
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

doc = Document()

# Set page margins to be small to fit 1 page
for section in doc.sections:
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 139) # Dark blue like the original
    
    # Add bottom border for the lines
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '00008B')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_contact_info():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run1 = p.add_run('VISHAL JHA\n')
    run1.bold = True
    run1.font.name = 'Calibri'
    run1.font.size = Pt(22)
    run1.font.color.rgb = docx.shared.RGBColor(0, 0, 139)
    run2 = p.add_run('+91-9672117679 | vishaljha572004@gmail.com | linkedin.com/in/vishaljha572004 | vishaljha572004-png | Jaipur, Rajasthan')
    run2.font.size = Pt(9.5)

def add_split_line(left_text, right_text, bold_left=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    # add tab stop for right alignment at 7.6 inches (compensating for 0.45 margin)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.6), WD_TAB_ALIGNMENT.RIGHT)
    
    run1 = p.add_run(left_text)
    if bold_left:
        run1.bold = True
    p.add_run('\t')
    run2 = p.add_run(right_text)

def add_paragraph(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    
def add_bullet(text, right_text=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0 
    if right_text:
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run(text)
        p.add_run('\t')
        p.add_run(right_text)
    else:
        p.add_run(text)
    
def add_tech_stack(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)

def add_skill(category, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run1 = p.add_run(category + ": ")
    run1.bold = True
    p.add_run(text)

add_contact_info()

add_heading('CAREER OBJECTIVE')
add_paragraph('Final-year B.Tech Computer Science student with hands-on experience in full-stack web development using React.js, Node.js, Express.js, MongoDB, and MySQL. Skilled in building scalable, responsive web applications, RESTful APIs, and integrating third party services with JWT-based authentication and data protection best practices. Passionate about technology and problem solving, eager to contribute as a Full Stack Developer in a collaborative, cross-functional team environment.')

add_heading('EXPERIENCE')
add_split_line('Full Stack Developer Intern – Cognifyz', 'May 2026 – June 2026')
add_bullet('Developed and integrated RESTful APIs using Node.js and Express.js for dynamic data handling, optimizing endpoint response structure and server-side logic.')
add_bullet('Built and maintained React-based UI components with efficient state management, improved UX by 30% across multiple application modules.')
add_bullet('Collaborated with cross-functional teams, contributing to production-ready features and participating in code reviews to maintain high code quality standards.')
add_bullet('Troubleshot, debugged, and optimized frontend rendering performance for smoother interactions on key user-facing pages.')

add_heading('PROJECTS')
add_split_line('V Mart – E-Commerce Grocery Website', 'May 2026 – Present')
add_tech_stack('Tech Stack: ReactJS, Node.js, Express.js, MongoDB, MySQL, JWT, REST APIs')
add_bullet('Designed and developed a responsive, user-friendly full-stack grocery e-commerce platform with 10+ product categories, cart management, and secure JWT-based authentication and data protection.')
add_bullet('Developed dynamic shopping cart with real-time price calculation, discount/coupon system, and seamless checkout flow with third-party API integration.')
add_bullet('Built admin dashboard managing 10+ product categories with inventory management, order tracking, and sales analytics — optimizing application performance and scalability.')

add_split_line('Talent Scope – AI-Powered Interview Preparation Platform', 'March 2026')
add_tech_stack('Tech Stack: ReactJS, Node.js, Express.js, MongoDB, Firebase, REST APIs, JWT, Razorpay')
add_bullet('Built a scalable AI-powered platform generating 5 role-specific questions per session, with voice-based responses via Web Speech API and animated AI avatar — ensuring seamless user experience across modules.')
add_bullet('Implemented ATS resume checker analyzing resumes against job descriptions with 95% keyword accuracy, plus PDF parsing for personalized question generation using third-party API integration.')
add_bullet('Integrated Razorpay payment gateway with credits-based premium access, JWT authentication, and automated PDF performance reports reducing manual evaluation time by 60%.')

add_heading('TECHNICAL SKILLS')
add_skill('Frontend', 'ReactJS, JavaScript (ES6+), HTML5, CSS3, Tailwind CSS')
add_skill('Backend', 'Node.js, Express.js, RESTful APIs, Server-side Logic')
add_skill('Databases', 'MySQL, MongoDB, Firebase, Redis')
add_skill('Tools & Technologies', 'Git, GitHub, VS Code, Postman, JWT, OAuth')
add_skill('Languages', 'C++, JavaScript')
add_skill('Concepts', 'OOP, DBMS, Web Security, Authentication & Authorization, REST APIs, Microservices')
add_skill('Soft Skills', 'Problem-Solving, Teamwork, Cross-functional Collaboration, Communication, Continuous Learning')

add_heading('EDUCATION')
add_split_line('B.Tech – Computer Science & Engineering', '2023 – 2027')
add_paragraph('JECRC University, Jaipur, Rajasthan | CGPA: 8.27 / 10.0')

add_heading('CERTIFICATE')
add_bullet('IBM Skills Network – SQL and Relational Databases 101 | IBM Cognitive Class', 'June 2026')
add_bullet('Goldman Sachs Software Engineering Job Simulation – Forage', 'Feb 2025')

# Make sure directory exists
output_dir = r'C:\Users\visha\Downloads\resume'
os.makedirs(output_dir, exist_ok=True)
doc.save(os.path.join(output_dir, 'Vishal_Jha_Resume_Updated.docx'))
print("Document created successfully at " + os.path.join(output_dir, 'Vishal_Jha_Resume_Updated.docx'))
